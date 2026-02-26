"""
Resume Exporter service — convert parsed resume analysis into
multiple output formats (PDF, DOCX, JSON, TXT).

Each export function accepts a Resume ORM model (with populated ``analysis``
JSON) and a base output path, returning the final file path.
"""

from __future__ import annotations

import json
import os
from typing import Any, Dict

import structlog

logger = structlog.get_logger(__name__)


def _get_analysis(resume) -> Dict[str, Any]:
    """Safely extract analysis dict from a Resume model."""
    analysis = getattr(resume, "analysis", None) or {}
    if not isinstance(analysis, dict):
        analysis = {}
    return analysis


# ── TXT Export ──────────────────────────────────────────────────────────────

async def export_to_txt(resume, base_path: str) -> str:
    """Export resume analysis as a formatted plain-text file."""
    analysis = _get_analysis(resume)
    file_path = f"{base_path}.txt"

    lines = [
        f"RESUME: {getattr(resume, 'title', 'Untitled')}",
        "=" * 60,
        "",
    ]

    # Summary
    summary = analysis.get("summary", "")
    if summary:
        lines += ["SUMMARY", "-" * 40, summary, ""]

    # Skills
    skills = analysis.get("skills", [])
    if skills:
        lines += ["SKILLS", "-" * 40]
        lines += [f"  • {s}" for s in skills]
        lines.append("")

    # Experience
    experience = analysis.get("experience", [])
    if experience:
        lines += ["WORK EXPERIENCE", "-" * 40]
        for exp in experience:
            company = exp.get("company", "Unknown")
            position = exp.get("position", "Unknown")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "Present")
            lines.append(f"  {position} at {company} ({start} – {end})")
            desc = exp.get("description", "")
            if desc:
                lines.append(f"    {desc}")
            lines.append("")

    # Education
    education = analysis.get("education", [])
    if education:
        lines += ["EDUCATION", "-" * 40]
        for edu in education:
            institution = edu.get("institution", "Unknown")
            degree = edu.get("degree", "")
            field = edu.get("field_of_study", "")
            lines.append(f"  {degree} in {field} — {institution}")
        lines.append("")

    # Recommendations
    recommendations = analysis.get("recommendations", [])
    if recommendations:
        lines += ["RECOMMENDATIONS", "-" * 40]
        lines += [f"  {i+1}. {r}" for i, r in enumerate(recommendations)]
        lines.append("")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info("resume_exported_txt", resume_id=str(resume.id))
    return file_path


# ── JSON Export ─────────────────────────────────────────────────────────────

async def export_to_json(resume, base_path: str) -> str:
    """Export resume analysis as a structured JSON file."""
    analysis = _get_analysis(resume)
    file_path = f"{base_path}.json"

    export_data = {
        "resume_id": str(resume.id),
        "title": getattr(resume, "title", "Untitled"),
        "inferred_role": getattr(resume, "inferred_role", None),
        "years_of_experience": getattr(resume, "years_of_experience", None),
        "confidence_score": getattr(resume, "confidence_score", None),
        "skills": analysis.get("skills", []),
        "experience": analysis.get("experience", []),
        "education": analysis.get("education", []),
        "summary": analysis.get("summary", ""),
        "recommendations": analysis.get("recommendations", []),
    }

    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(export_data, f, indent=2, default=str)

    logger.info("resume_exported_json", resume_id=str(resume.id))
    return file_path


# ── PDF Export ──────────────────────────────────────────────────────────────

async def export_to_pdf(resume, base_path: str) -> str:
    """Export resume analysis as a PDF file.

    Uses ``reportlab`` if available; otherwise falls back to a text-based
    approach wrapped in a minimal PDF structure.
    """
    file_path = f"{base_path}.pdf"

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        analysis = _get_analysis(resume)
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(getattr(resume, "title", "Resume"), styles["Title"]))
        story.append(Spacer(1, 12))

        # Summary
        summary = analysis.get("summary", "")
        if summary:
            story.append(Paragraph("Summary", styles["Heading2"]))
            story.append(Paragraph(summary, styles["Normal"]))
            story.append(Spacer(1, 12))

        # Skills
        skills = analysis.get("skills", [])
        if skills:
            story.append(Paragraph("Skills", styles["Heading2"]))
            story.append(Paragraph(", ".join(skills), styles["Normal"]))
            story.append(Spacer(1, 12))

        # Experience
        experience = analysis.get("experience", [])
        if experience:
            story.append(Paragraph("Work Experience", styles["Heading2"]))
            for exp in experience:
                company = exp.get("company", "")
                position = exp.get("position", "")
                story.append(Paragraph(f"<b>{position}</b> at {company}", styles["Normal"]))
                desc = exp.get("description", "")
                if desc:
                    story.append(Paragraph(desc, styles["Normal"]))
                story.append(Spacer(1, 6))

        # Education
        education = analysis.get("education", [])
        if education:
            story.append(Paragraph("Education", styles["Heading2"]))
            for edu in education:
                institution = edu.get("institution", "")
                degree = edu.get("degree", "")
                field = edu.get("field_of_study", "")
                story.append(Paragraph(f"<b>{degree}</b> in {field} — {institution}", styles["Normal"]))
                story.append(Spacer(1, 6))

        doc.build(story)

    except ImportError:
        # Fallback: generate a plain-text file renamed as .pdf
        # (not ideal — install reportlab for real PDF output)
        logger.warning("reportlab_not_installed_falling_back_to_txt_pdf")
        txt_path = await export_to_txt(resume, base_path + "_pdf_fallback")
        os.rename(txt_path, file_path)

    logger.info("resume_exported_pdf", resume_id=str(resume.id))
    return file_path


# ── DOCX Export ─────────────────────────────────────────────────────────────

async def export_to_docx(resume, base_path: str) -> str:
    """Export resume analysis as a DOCX file.

    Requires ``python-docx`` (already in requirements for upload parsing).
    """
    file_path = f"{base_path}.docx"
    analysis = _get_analysis(resume)

    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(getattr(resume, "title", "Resume"), level=0)

        # Summary
        summary = analysis.get("summary", "")
        if summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)

        # Skills
        skills = analysis.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        # Experience
        experience = analysis.get("experience", [])
        if experience:
            doc.add_heading("Work Experience", level=1)
            for exp in experience:
                company = exp.get("company", "")
                position = exp.get("position", "")
                p = doc.add_paragraph()
                run = p.add_run(f"{position} at {company}")
                run.bold = True
                desc = exp.get("description", "")
                if desc:
                    doc.add_paragraph(desc, style="List Bullet")

        # Education
        education = analysis.get("education", [])
        if education:
            doc.add_heading("Education", level=1)
            for edu in education:
                institution = edu.get("institution", "")
                degree = edu.get("degree", "")
                field = edu.get("field_of_study", "")
                doc.add_paragraph(f"{degree} in {field} — {institution}")

        # Recommendations
        recommendations = analysis.get("recommendations", [])
        if recommendations:
            doc.add_heading("Recommendations", level=1)
            for rec in recommendations:
                doc.add_paragraph(rec, style="List Bullet")

        doc.save(file_path)

    except ImportError:
        logger.warning("python_docx_not_installed_falling_back_to_txt")
        txt_path = await export_to_txt(resume, base_path + "_docx_fallback")
        os.rename(txt_path, file_path)

    logger.info("resume_exported_docx", resume_id=str(resume.id))
    return file_path
